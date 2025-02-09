import pandas as pd
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import networkx as nx
from math import cos, sin, pi
from docx import Document
from docx.shared import Inches

# ----------------------
# 1. Bar Graph
# ----------------------

def create_bar_chart():
    # Load and preprocess the data
    df = pd.read_csv('bar_assignment.csv')
    df.replace({0: 'No', 1: 'Yes'}, inplace=True)
    df_grouped = df.groupby(['LABEL', 'COUNT']).size().unstack(fill_value=0)

    # Create the plot
    fig, ax = plt.subplots(figsize=(10, 6))
    bar_plot = df_grouped.plot(
        kind='barh', 
        stacked=True, 
        ax=ax, 
        color=['red', 'blue']  # Use distinct colors for "Yes" and "No"
    )

    # Set title and axis labels
    ax.set_title('Bar Graph', loc='left', fontsize=14, pad=20)
    ax.set_xlabel('Count', fontsize=12)
    ax.set_ylabel('Label', fontsize=12)

    # Move legend below the title, neatly positioned
    ax.legend(
        title='LEGEND', 
        loc='upper right', 
        bbox_to_anchor=(0, 1.02), 
        ncol=2, 
        frameon=False, 
        fontsize=10
    )

    # Add numbers to bars, slightly inside the bars aligned to the left
    for container in ax.containers:
        labels = [f'{int(value)}' if value > 0 else '' for value in container.datavalues]
        ax.bar_label(container, labels=labels, label_type='edge', padding=3, fontsize=9)

    # Final layout adjustments
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    plt.savefig('bar_chart.png', dpi=300, bbox_inches='tight')
    plt.close()

# ----------------------
# 2. Sankey Diagram
# ----------------------

def create_sankey():
    df = pd.read_csv('sankey_assignment.csv')
    
    # Define nodes in three sections: Left, Middle, Right
    left_nodes = ['OMP', 'PS', 'CNP', 'NCDM', 'RGS', 'NRP', 'PEC', 'NMCCC']
    middle_nodes = ['I', 'S', 'D', 'N', 'F']
    right_nodes = ['Aca', 'Oth', 'Reg']
    
    # Combine all nodes into a single list
    nodes = left_nodes + middle_nodes + right_nodes
    node_indices = {node: i for i, node in enumerate(nodes)}
    
    # Create links
    source = []
    target = []
    value = []
    
    # Links from Left to Middle
    for _, row in df.iterrows():
        label = row['LABEL']  # Middle node (I, S, D, N, F)
        for col in left_nodes:
            if row[col] > 0:
                source.append(node_indices[col])  # Left node
                target.append(node_indices[label])  # Middle node
                value.append(row[col])
    
    # Links from Middle to Right
    # Assuming the connections are based on the LABEL column
    for _, row in df.iterrows():
        label = row['LABEL']  # Middle node (I, S, D, N, F)
        for col in right_nodes:
            if row[col] > 0:
                source.append(node_indices[label])  # Middle node
                target.append(node_indices[col])  # Right node
                value.append(row[col])
    
    # Create Sankey diagram
    fig = go.Figure(go.Sankey(
        node=dict(
            label=nodes,
            color=['blue'] * len(left_nodes) + ['green'] * len(middle_nodes) + ['orange'] * len(right_nodes)
        ),
        link=dict(source=source, target=target, value=value)
    ))
    fig.update_layout(title_text="Sankey Diagram", font_size=10)
    fig.write_image("sankey.png", scale=2)

# ----------------------
# 3. Network Graph
# ----------------------
def create_network():
    df = pd.read_csv('networks_assignment.csv')
    G = nx.Graph()
    
    # Add edges
    for _, row in df.iterrows():
        source = row['LABELS']
        for target, weight in row[1:].items():
            if weight > 0:
                G.add_edge(source, target, weight=weight)
    
    # Define node colors
    pentagon = ['D', 'F', 'I', 'N', 'S']
    green = ['BIH', 'GEO', 'ISR', 'MNE', 'SRB', 'CHE', 'TUR', 'UKR', 'GBR', 'AUS', 'HKG', 'USA']
    yellow = ['AUT', 'BEL', 'BGR', 'HRV', 'CZE', 'EST', 'FRA', 'DEU', 'GRC', 'HUN', 'IRL', 'ITA', 'LVA', 'LUX', 'NLD', 'PRT', 'ROU', 'SVK', 'SVN', 'ESP']
    
    # Assign colors to all nodes
    node_colors = []
    for node in G.nodes:
        if node in pentagon:
            node_colors.append('blue')
        elif node in green:
            node_colors.append('green')
        elif node in yellow:
            node_colors.append('yellow')
        else:
            node_colors.append('red')  # Default color for any unclassified nodes
    
    # Position nodes
    pos = {}
    # Place pentagon nodes in a star
    angles = [pi/2 + 2*pi*i/5 for i in range(5)]  # 5 angles for the star
    for i, node in enumerate(pentagon):
        pos[node] = (cos(angles[i]), sin(angles[i]))
    # Place other nodes in a circle
    other_nodes = [n for n in G.nodes if n not in pentagon]
    other_pos = nx.circular_layout(other_nodes, scale=2)
    pos.update(other_pos)
    
    plt.figure(figsize=(10, 10))
    nx.draw(G, pos, with_labels=True, node_color=node_colors, edge_color='gray')
    
    # Add a label to the network graph
    plt.text(
        0.5, 1.05, 'Network Graph', 
        transform=plt.gca().transAxes, 
        fontsize=12, 
        ha='center', 
        va='center'
    )
    
    plt.savefig('network.png', dpi=300, bbox_inches='tight')
    plt.close()

# ----------------------
# 4. Collate Graphs
# ----------------------
def collate_graphs():
    fig, axes = plt.subplots(3, 1, figsize=(8.5, 11))  # Long bond paper size (8.5x11 inches)
    
    # Load images
    bar_img = plt.imread('bar_chart.png')
    sankey_img = plt.imread('sankey.png')
    network_img = plt.imread('network.png')
    
    # Display images
    axes[0].imshow(bar_img)
    axes[0].axis('off')
    
    axes[1].imshow(sankey_img)
    axes[1].axis('off')
    
    axes[2].imshow(network_img)
    axes[2].axis('off')
    
    plt.tight_layout()
    plt.savefig('collated.png', dpi=300, bbox_inches='tight')
    plt.close()

# ----------------------
# 5. Create PDF Sample
# ----------------------
def create_pdf():
    doc = Document()
    doc.add_heading('Collated Graphs', 0)
    doc.add_picture('collated.png', width=Inches(8.5))
    doc.save('collated_sample.docx')
    # Manually save as PDF from Word

# ----------------------
# Run All Functions
# ----------------------
if __name__ == "__main__":
    create_bar_chart()
    create_sankey()
    create_network()
    collate_graphs()
    create_pdf()